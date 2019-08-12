import zipfile
from zipfile import ZipFile
import os
from os import path as osp
import shutil
import pdb
from PIL import Image, ImageFilter
import io
import re
import xml.etree.ElementTree as ET
from docx import Document
import argparse
import sys
#from docx.shared import Inches

################################################################################
# ReportMod.Base
#   Functions and Variables used in other classes will be defined here
#   Current Variables:
#       Pictures - Extensions of Picture Files
#       wdir - Current Working Directory
#       blur - Default Blur settting
#       PictureSizes - Picture Height in Pixels, to categorize pictures on sizes
#   Current Functions
#       Blur() - Returns a blur value based on Picture size
#       Shrink() - Returns size of image thumbnail based on original size
#       Unzip() - Takes a file, unzips it into a new directory
#       SizeImage() - Takes a picture and determines Size
#       createNG() - Creates a Regex with a Named Group
#       NamedGroup() - Returns matches for a Named Group
#       FindAll() - Returns all matches for a Regex. Handles text search as well
################################################################################
class Base:
    pictures = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']
    wdir = os.getcwd()
    blur = ImageFilter.GaussianBlur(40)
    PictureSizes = {
        "Locator": 1,
        "Icon" : 64,
        "Tiny" : 100,
        "Small": 150,
        "Decent": 400,
        "Medium": 800,
        "Big": 1250,
        "Large": 2000,
        "Giant": 4000
    }
    fillSequence = '****'

    @staticmethod
    def cd(folder):
        Base.wdir = folder
        os.chdir(folder)

    @staticmethod
    def Blur(img):
        '''
        This function dynamically blurs images based on size In larger images,
            this should ease cases where larger pics look unchanged 
        '''
        size = Base.SizeImage(img)
        blurs = {
            "Locator": 0,
            "Icon" : 20,
            "Tiny" : 40,
            "Small": 40,
            "Decent": 40,
            "Medium": 40,
            "Big": 40,
            "Large": 80,
            "Giant": 100
        }
        return ImageFilter.GaussianBlur(blurs[size])

    @staticmethod
    def Shrink(img):
        '''
        This function shrinks the picture according to original image size
        '''
        newSize = {
            "Locator": 1,
            "Icon" : 32,
            "Tiny" : 32,
            "Small": 32,
            "Decent": 32,
            "Medium": 64,
            "Big": 64,
            "Large": 128,
            "Giant": 128
        }
        oldSize = Base.SizeImage(img)
        if oldSize == "Locator":
            return 1,1
        else:
            return (newSize[oldSize], newSize[oldSize])

    @staticmethod
    def Unzip(file):
        '''
        Docx files can be operated as zip files.
        '''
        with ZipFile(file, 'r') as ziph:
            os.chdir(Base.wdir)
            os.mkdir(osp.splitext(osp.basename(file))[0])
            os.chdir(osp.splitext(osp.basename(file))[0])
            ziph.extractall()
            os.chdir(Base.wdir)
        return osp.join(Base.wdir, osp.splitext(osp.basename(file))[0])

    @staticmethod
    def SizeImage(img):
        '''
        This function will take an image, find it's height, and guage which
            category it should be placed in.
        '''
        match = ret = False
        with open(img, 'r') as fh:
            x,y = Image.open(img).size
            for elem in Base.PictureSizes.items():
                match = y > elem[1]
                if match == True and elem[0] != "Giant":
                    continue
                else:
                    ret = elem[0]
                    break
        return ret

    @staticmethod
    def createNG(namedgroup, ngRGX, PreRGX=None, PostRGX=None):
        namedgroup = namedgroup.replace(' ','').strip(',.;:?"\'').lower()
        rgx = PreRGX + r'(P?<' + namedgroup + r'>' + ngRGX + r')' + PostRGX
        return rgx

    @staticmethod
    def NamedGroup(namedgroup, regex, corpus, multiline=False):
        re.MULTILINE = multiline
        if r'(?P<' in regex == False:
            return None, None
        match = re.search(regex, corpus)
        if match:
            return namedgroup, match
        return None, None
        

    @staticmethod
    def FindAll(regex, corpus, multiline=False):
        re.MULTILINE = multiline
        match = re.findall(regex, corpus)
        if match:
            return match
        return None

################################################################################
# ReportMod.Docx
#   This module is meant to work specifically with the Docx files
#   Please Note that Doc files ARE NOT Docx Files
#   Current Functions
#       RemoveObjects() - Removes files from Docx Archive
#       BlurImages() - Takes every image from the docs archive and blur it
#       RedactRegex() - Takes a regex, and redacts it in Word
################################################################################
class Docx:
    '''
    Explicitly for Docx only, not for doc files.
    Open Document Text files (odt) currently work, but isn't the intended function of this class
        I plan to upkeep this as long as odt doesn't require too much extra work
        At which point, it'll have its own class
    '''
    baseschema = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    paragraph = baseschema + 'p'
    text = baseschema + 't'

    @staticmethod
    def RemoveObjects(zfile, ext=[]):
        '''
        This function requires docx files to already be unzipped. 
        This function will cause read errors.
        Consider this when you need to remove pictures or attachements
            But don't want to endanger formatting
        Unfortunately, current Python libraries do not currently support
            removing objects cleanly. Implementing this feature sans
            read-errors is already on the feature list for this library.
        '''
        zfolder = Base.Unzip(zfile)
        os.chdir(Base.wdir)
        ziph = ZipFile(zfolder + '.redacted.docx', 'w', zipfile.ZIP_DEFLATED)
        for root, dirs, files in os.walk(zfolder):
            for file in files:
                # THANK YOU!!! to Martijn Pieters for his help!
                # https://stackoverflow.com/questions/35937121/python-zip-folder-without-including-current-directory
                path = os.path.normpath(os.path.join(root, file))
                if osp.splitext(file)[1] not in ext and os.path.isfile(path):
                    ziph.write(path, osp.relpath(path, osp.abspath(zfolder)))
        ziph.close()

    @staticmethod
    def BlurImages(ifile, ofile=None):
        '''
        If blurring is good enough, this function will take the document and
        blur the image.
        To change the amount of blurring, change Docx.blur
            The Default value is: ImageFilter.GaussianBlur(50)
            Adjust the Number inside GaussianBlur until the valus suits you
        '''
        # Thanks to mata for his/her help
        # https://stackoverflow.com/questions/50963852/remove-all-images-from-docx-files
        if ofile is None and osp.splitext(ifile)[1] == '.docx':
            ofile = ifile.replace('.docx','.redacted.docx')
        elif ofile is None and osp.splitext(ifile)[1] == '.odt':
            ofile = ifile.replace('.odt','.redacted.odt')
        elif ofile is None:
            ofile = osp.splitext(ifile)[0] + '.redacted.' \
                    + osp.splitext(ifile)[1]
        izip = ZipFile(ifile)
        ozip = ZipFile(ofile, 'w', zipfile.ZIP_DEFLATED)
        for node in izip.infolist():
            content = izip.read(node)
            if osp.splitext(node.filename)[1] in Base.pictures:
                # Blur Image and save it to a buffer
                img = Image.open(io.BytesIO(content)) # Load Picture to memory
                img = img.convert().filter(Base.blur) # Blur it
                imgBuffer = io.BytesIO()
                img.save(imgBuffer, osp.splitext(node.filename)[1][1:])
                # Change the XML so that it points to the blurred image, not original
                content = imgBuffer.getvalue()
                node.file_size = len(content)
                node.CRC = zipfile.crc32(content)
                img.close()
            ozip.writestr(node, content)
        izip.close()
        ozip.close()

    @staticmethod
    def RedactRegex(ifile, regex, redactText=Base.fillSequence, ofile=None):
        reg = re.compile(regex)
        doc = Document(ifile)
        if ofile is None:
            ofile = osp.splitext(ifile)[0] + '.redacted' \
                    + osp.splitext(ifile)[1]
        # Search in Paragraphs (normal text in Word)
        for para in doc.paragraphs:
            if reg.search(para.text):
                para.text = reg.sub(redactText, para.text)
        # Search in Tables (A separate entity altogether)
        for table in doc.tables:
            for cell in table.cells:
                for para in cell.paragraphs:
                    if reg.search(para.text):
                        para.text = reg.sub(redactText, para.text)
        doc.save(ofile)

    @staticmethod
    def redactRegex2(wfile, regex, redactText=Base.fillSequence):
        '''
        WORK IN PROGRESS, USE Docx.redactRegex INSTEAD
        '''
        regex = re.compile(regex)
        wordfolder = Base.Unzip(wfile)
        docpath = osp.abspath(wordfolder) + r'\word\document.xml'
        with open(docpath, 'r+') as docpathh:
            content = docpathh.read()
            xmlc = ET.XML(content)
            # Adapted from https://etienned.github.io/posts/extract-text-from-word-docx-simply/
            for i in xmlc.getiterator(Docx.paragraph):
                for j in i.getiterator(Docx.text):
                    if j:
                        text = regex.sub(redactText, text)
                        j.text = text
            docpathh.write(ET.tostring(xmlc, encoding='utf-8', method='xml').decode())

################################################################################
# ReportMod.HTML
#   This module works on HTML reports, which are very easy to work with
#   Current Functions:
#       SetReports() - Duplicates a HTML report, so that it isn't affected
#       BlurImages() - Takes images in the report and blurs them
#       ShrinkImages() - Takes images in the report and shrinks them
################################################################################
class HTML:
    # Have to leave out js/css files, as you might modify code!!
    # Add rtf here
    textFiles = ['.html', '.xml', '.txt', '.csv', '.tsv', '.json']
    codeFiles = ['.js', '.css'] # Add these at your risk
    
    @staticmethod
    def SetReports(ifolder, ofolder):
        try:
            shutil.copytree(ifolder, ofolder, copy_function=shutil.copy2)
        except shutil.Error as e:
            print("Error: " + e)
        except OSError as e:
            print("Error: " + e)
        Base.cd(ofolder)

    @staticmethod
    def BlurImages():
        for root, dirs, files in os.walk(Base.wdir):
            for file in files:
                if osp.splitext(file)[1] in Base.pictures:
                    blursize = Base.Blur(osp.join(root, file))
                    fh = open(osp.join(root, file), 'rb')
                    img = Image.open(io.BytesIO(fh.read()))
                    img = img.convert().filter(blursize)
                    img.save(osp.join(root, file))
                    img.close()
                    fh.close()

    @staticmethod
    def ShrinkImages():
        for root, dirs, files in os.walk(Base.wdir):
            for file in files:
                if osp.splitext(file)[1].lower() in Base.pictures:
                    imgsize = Base.Shrink(osp.join(root, file))
                    fh = open(osp.join(root, file), 'rb')
                    img = Image.open(io.BytesIO(fh.read()))
                    if imgsize[0] == 1:
                        continue
                    else:
                        img.thumbnail(imgsize, Image.ANTIALIAS)
                    try:
                        img.save(osp.join(root, file))
                    except IOError:
                        print("Unable to Create File")
                    img.close()
                    fh.close()

    @staticmethod
    def RedactRegex(regex, redactText=Base.fillSequence):
        reg = re.compile(regex)
        for root, firs, files in os.walk(Base.wdir):
            for file in files:
                if osp.splitext(file)[1] in HTML.textFiles:
                    with open(osp.join(root, file), 'r+') as fileh:
                        filep = fileh.tell()
                        line = fileh.readline()
                        while line:
                            if reg.search(line):
                                line = reg.sub(redactText, line)
                                fileh.seek(filep)
                                fileh.write(line)
                            filep = fileh.tell()
                            line = fileh.readline()

################################################################################
# Autorun:
#   1) Default Values - Minimum values to run
#   2) Loads default Config Files if available
#   3) Loads Custom Config Files if available
#   4) Should take argparse command-line arguments
#   5) Cleans up any leftover directory hang ups
################################################################################
if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='ReportMod POC - Report Modifier Tool to change reports on the fly')
    DocFileType = parser.add_mutually_exclusive_group(required=True)
    DocFileType.add_argument('--html', action='store_true', help='Input html folder – Necessary if docx is not chosen')    
    DocFileType.add_argument('--docx', action='store_true', help='Input docx file – Necessary if html is not chosen')    
    parser.add_argument('-i', dest='InputReport', required=True, help='The name of the initial report - necessary') 
    parser.add_argument('-o', dest='OutputReport', required=True, help='Name of output folder or file – necessary')
    parser.add_argument('-e', dest='RegexString', help='Regex String to search for and replace - Optional')
    parser.add_argument('-s', dest='RedactionString', help='String to replace matches - Default is ***** - Optional')
    parser.add_argument('-b', action='store_true', help='Blur all images of the report')
    parser.add_argument('-t', action='store_true', help='Shrink all images of the report - HTML only right now')
    args = parser.parse_args()

    if osp.exists(args.OutputReport):
        print("\nOutput Report already exists. Will not overwrite file")
        sys.exit()

    if args.html:
        HTML.SetReports(args.InputReport, args.OutputReport)
        if args.RegexString:
            if args.RedactionString:
                HTML.RedactRegex(args.RegexString, args.RedactionString)
            else:
                HTML.RedactRegex(args.RegexString)
        if args.t:
            HTML.ShrinkImages()
        if args.b:
            HTML.BlurImages()
    elif args.docx:
        Base.cd(osp.abspath(osp.join(args.OutputReport, '..')))
        if args.RegexString:
            if args.b:
                tmpReport = osp.splitext(args.OutputReport)[0] + '.blur' \
                                + osp.splitext(args.OutputReport)[1]
                Docx.BlurImages(args.InputReport, tmpReport)
                if args.RedactionString:
                    Docx.RedactRegex(tmpReport, args.RegexString,
                                     args.RedactionString, args.OutputReport)
                    os.remove(tmpReport)
                    sys.exit()
                else:
                    Docx.RedactRegex(tmpReport, args.RegexString,
                                     ofile=args.OutputReport)
                    os.remove(tmpReport)
                    sys.exit()
            else:
                if args.RedactionString:
                    Docx.RedactRegex(args.InputReport, args.RegexString,
                                     args.RedactionString, args.OutputReport)
                else:
                    Docx.RedactRegex(args.InputReport, args.RegexString,
                                     ofile=args.OutputReport)
        if args.t:
            print('\nProof Of Concept Limit: No Shrinking exists yet for Docx')
        else:
            if args.b:
                Docx.BlurImages(args.InputReport, args.OutputReport)
